import io
import re
import math
import unicodedata
from difflib import SequenceMatcher

import pandas as pd


# ==========================================================
# VŠEOBECNÉ POMOCNÉ FUNKCIE
# ==========================================================

def _normalize_text(value):
    if value is None:
        return ""

    text = str(value).strip().lower()

    text = unicodedata.normalize(
        "NFKD",
        text
    )

    text = "".join(
        char
        for char in text
        if not unicodedata.combining(char)
    )

    text = text.replace("×", "x")
    text = text.replace("–", "-")
    text = text.replace("—", "-")

    text = re.sub(
        r"\s+",
        " ",
        text
    ).strip()

    return text


def _normalize_code(value):
    """
    Zjednotí napr.:
    631362442.S
    631362442.s
    631362442
    """

    text = _normalize_text(
        value
    )

    text = re.sub(
        r"\.s$",
        "",
        text
    )

    return text


def _to_number(value):
    if value is None:
        return None

    if isinstance(
        value,
        bool
    ):
        return None

    if isinstance(
        value,
        (int, float)
    ):
        if isinstance(
            value,
            float
        ) and math.isnan(value):
            return None

        return float(value)

    text = str(
        value
    ).strip()

    if not text:
        return None

    text = (
        text
        .replace("\xa0", "")
        .replace(" ", "")
        .replace(",", ".")
    )

    try:
        return float(text)

    except Exception:
        return None


def _normalize_unit(value):
    text = _normalize_text(
        value
    )

    unit_map = {
        "m²": "m2",
        "m2": "m2",
        "m³": "m3",
        "m3": "m3",
        "ks": "ks",
        "kus": "ks",
        "kusy": "ks",
        "kompl": "kompl",
        "komplet": "kompl",
        "súbor": "subor",
        "subor": "subor",
    }

    return unit_map.get(
        text,
        text
    )


# ==========================================================
# ČÍTANIE PDF / DOCX / EXCEL PRE OSTATNÉ ČASTI APPKY
# ==========================================================

def extract_text_from_pdf(file_bytes):
    try:
        from pypdf import PdfReader
    except ImportError:
        from PyPDF2 import PdfReader

    reader = PdfReader(
        io.BytesIO(file_bytes)
    )

    parts = []

    for page_number, page in enumerate(
        reader.pages,
        start=1
    ):
        text = (
            page.extract_text()
            or ""
        )

        parts.append(
            f"\n--- STRANA {page_number} ---\n{text}"
        )

    return "\n".join(
        parts
    )


def extract_text_from_docx(file_bytes):
    from docx import Document

    document = Document(
        io.BytesIO(file_bytes)
    )

    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            parts.append(
                text
            )

    for table_number, table in enumerate(
        document.tables,
        start=1
    ):
        parts.append(
            f"\n--- TABUĽKA {table_number} ---"
        )

        for row in table.rows:
            values = [
                cell.text.strip()
                for cell in row.cells
            ]

            if any(values):
                parts.append(
                    " | ".join(values)
                )

    return "\n".join(
        parts
    )


def extract_excel_rows(file_bytes):
    sheets = pd.read_excel(
        io.BytesIO(file_bytes),
        sheet_name=None,
        header=None,
        dtype=object
    )

    result = []

    for sheet_name, dataframe in sheets.items():

        for dataframe_index, row in dataframe.iterrows():

            values = []

            for value in row.tolist():

                if pd.isna(value):
                    values.append("")
                else:
                    values.append(
                        str(value).strip()
                    )

            if not any(
                value != ""
                for value in values
            ):
                continue

            result.append(
                {
                    "sheet": str(
                        sheet_name
                    ),
                    "row_number": int(
                        dataframe_index
                    ) + 1,
                    "values": values
                }
            )

    return result


def extract_text_from_excel(file_bytes):
    rows = extract_excel_rows(
        file_bytes
    )

    if not rows:
        return ""

    parts = []
    current_sheet = None

    for row in rows:

        sheet = row.get(
            "sheet",
            ""
        )

        if sheet != current_sheet:
            current_sheet = sheet

            parts.append(
                f"\n--- LIST: {sheet} ---"
            )

        parts.append(
            "RIADOK "
            + str(
                row.get(
                    "row_number",
                    ""
                )
            )
            + ": "
            + " | ".join(
                row.get(
                    "values",
                    []
                )
            )
        )

    return "\n".join(
        parts
    )


# ==========================================================
# 1. VŠEOBECNÉ NAČÍTANIE ROZPOČTU
# ==========================================================

def extract_budget_items_python(file_bytes):
    """
    VŠEOBECNÁ funkcia.

    V každom hárku hľadá tabuľku s hlavičkami:
    Kód | Popis | MJ | Množstvo

    Nezávisí od názvu projektu ani od typu stavby.

    Rekapitulácie preskočí, aby sa položky
    nespočítali druhýkrát.

    NEVOLÁ AI.
    """

    sheets = pd.read_excel(
        io.BytesIO(file_bytes),
        sheet_name=None,
        header=None,
        dtype=object
    )

    items = []

    for sheet_name, dataframe in sheets.items():

        normalized_sheet_name = (
            _normalize_text(
                sheet_name
            )
        )

        # Rekapitulačný hárok zvyčajne obsahuje už
        # súčty detailných hárkov.
        if "rekapitul" in normalized_sheet_name:
            continue

        header_row_index = None
        header_columns = None

        for dataframe_index, row in dataframe.iterrows():

            normalized_values = [
                _normalize_text(value)
                for value in row.tolist()
            ]

            positions = {}

            aliases = {
                "kod": {
                    "kod",
                    "kód"
                },
                "popis": {
                    "popis",
                    "nazov",
                    "názov",
                    "popis polozky",
                    "nazov polozky"
                },
                "mj": {
                    "mj",
                    "m.j.",
                    "merna jednotka",
                    "merná jednotka"
                },
                "mnozstvo": {
                    "mnozstvo",
                    "množstvo"
                }
            }

            for target, accepted in aliases.items():

                for index, cell_text in enumerate(
                    normalized_values
                ):
                    if cell_text in {
                        _normalize_text(x)
                        for x in accepted
                    }:
                        positions[target] = index
                        break

            if len(positions) == 4:
                header_row_index = int(
                    dataframe_index
                )
                header_columns = positions
                break

        if (
            header_row_index is None
            or header_columns is None
        ):
            continue

        for dataframe_index in range(
            header_row_index + 1,
            len(dataframe)
        ):

            row = dataframe.iloc[
                dataframe_index
            ]

            code = row.iloc[
                header_columns["kod"]
            ]

            description = row.iloc[
                header_columns["popis"]
            ]

            unit = row.iloc[
                header_columns["mj"]
            ]

            quantity_raw = row.iloc[
                header_columns["mnozstvo"]
            ]

            quantity = _to_number(
                quantity_raw
            )

            if quantity is None:
                continue

            if (
                pd.isna(code)
                or pd.isna(description)
                or pd.isna(unit)
            ):
                continue

            code = str(
                code
            ).strip()

            description = str(
                description
            ).strip()

            unit = str(
                unit
            ).strip()

            if (
                not code
                or not description
                or not unit
            ):
                continue

            items.append(
                {
                    "sheet": str(
                        sheet_name
                    ),
                    "row_number": (
                        int(
                            dataframe_index
                        )
                        + 1
                    ),
                    "code": code,
                    "description": description,
                    "unit": _normalize_unit(
                        unit
                    ),
                    "quantity": quantity
                }
            )

    return items


# ==========================================================
# 2. TECHNICKÝ PODPIS POLOŽKY
# ==========================================================

def _remove_pricing_bands(text):
    """
    Odstraňuje iba typické CENOVÉ PÁSMA,
    ktoré nemenia technický význam práce.

    Príklady:
    - do 100 m3
    - nad 100 do 1000 m3
    - od 100 do 1000 m3
    - na vzdialenosť do 1000 m
    - za každých ďalších 1000 m

    Neodstraňuje DN, hrúbku, triedu betónu,
    materiál, SN, PN, rozmer výrobku atď.
    """

    result = _normalize_text(
        text
    )

    patterns = [
        # objemové / plošné / hmotnostné pásma
        r"\bnad\s+\d+(?:[.,]\d+)?\s+do\s+\d+(?:[.,]\d+)?\s*(?:m3|m2|m|t|kg|ks)\b",
        r"\bod\s+\d+(?:[.,]\d+)?\s+do\s+\d+(?:[.,]\d+)?\s*(?:m3|m2|m|t|kg|ks)\b",
        r"\bdo\s+\d+(?:[.,]\d+)?\s*(?:m3|m2|t|kg|ks)\b",
        r"\bnad\s+\d+(?:[.,]\d+)?\s*(?:m3|m2|t|kg|ks)\b",

        # vzdialenostné cenové pásma
        r"\bna vzdialenost do\s+\d+(?:[.,]\d+)?\s*m\b",
        r"\bna vzdialenost nad\s+\d+(?:[.,]\d+)?\s*do\s+\d+(?:[.,]\d+)?\s*m\b",
        r"\bza kazdych dalsich a zacatych\s+\d+(?:[.,]\d+)?\s*m\b",

        # plocha pracovného pruhu ako cenový interval
        r"\bplochy do\s+\d+(?:[.,]\d+)?\s*m2\b",
        r"\bplochy nad\s+\d+(?:[.,]\d+)?\s*do\s+\d+(?:[.,]\d+)?\s*m2\b",
    ]

    for pattern in patterns:
        result = re.sub(
            pattern,
            "",
            result,
            flags=re.IGNORECASE
        )

    result = re.sub(
        r"\s+",
        " ",
        result
    ).strip(" ,;-")

    return result


def _extract_critical_parameters(description):
    """
    Všeobecný technický podpis.

    Zachováva parametre, ktoré typicky znamenajú,
    že položky sa NESMÚ zlúčiť.
    """

    text = _normalize_text(
        description
    )

    parameters = []

    regexes = [
        # DN
        (
            "dn",
            r"\bdn\s*([0-9]+)\b"
        ),

        # SN
        (
            "sn",
            r"\bsn\s*([0-9]+)\b"
        ),

        # PN
        (
            "pn",
            r"\bpn\s*([0-9]+(?:[.,][0-9]+)?)\b"
        ),

        # trieda betónu C 20/25
        (
            "beton",
            r"\bc\s*([0-9]+)\s*/\s*([0-9]+)\b"
        ),

        # hrúbka
        (
            "hr",
            r"\bhr\.?\s*([0-9]+(?:[.,][0-9]+)?)\s*(mm|cm|m)\b"
        ),

        # priemer
        (
            "priemer",
            r"\bpriemer(?:u)?\s*([0-9]+(?:[.,][0-9]+)?)\s*(mm|cm|m)?\b"
        ),

        # rozmery 150x150, 1500x1800x2300
        (
            "rozmer",
            r"\b([0-9]+(?:[.,][0-9]+)?x[0-9]+(?:[.,][0-9]+)?(?:x[0-9]+(?:[.,][0-9]+)?)?)\s*(mm|cm|m)?\b"
        ),

        # pevnostné / triedové označenia bežné pri materiáloch
        (
            "xc",
            r"\b(xc[0-9]+|xf[0-9]+|xd[0-9]+|xa[0-9]+)\b"
        ),
    ]

    for name, pattern in regexes:

        for match in re.finditer(
            pattern,
            text,
            flags=re.IGNORECASE
        ):
            value = "|".join(
                part
                for part in match.groups()
                if part is not None
            )

            parameters.append(
                f"{name}:{value}"
            )

    # Materiálové a výrobkové tokeny.
    # Toto nie je zoznam stavieb; iba technické slová,
    # ktoré nesmú zmiznúť pri porovnaní.
    material_tokens = [
        "pvc-u",
        "pvc",
        "hdpe",
        "pe100",
        "pe80",
        "pp",
        "beton",
        "zelezo",
        "ocel",
        "nerez",
        "kari",
        "eps",
        "xps",
        "mineralna vlna",
        "sklenena vlna",
        "tehla",
        "porotherm",
        "ytong",
        "sadrokarton",
        "asfalt",
        "strkopiesok",
        "piesok",
        "kamenivo",
        "makadam",
        "drevo",
        "hlinik",
        "med",
    ]

    for token in material_tokens:

        if token in text:
            parameters.append(
                f"mat:{token}"
            )

    return tuple(
        sorted(
            set(
                parameters
            )
        )
    )


def _description_signature(description):
    """
    Vráti textový podpis po odstránení iba
    oceňovacích pásiem.
    """

    text = _remove_pricing_bands(
        description
    )

    # drobné typografické rozdiely
    text = re.sub(
        r"[(),.;:]",
        " ",
        text
    )

    text = re.sub(
        r"\s+",
        " ",
        text
    ).strip()

    return text


def _description_similarity(
    first,
    second
):
    return SequenceMatcher(
        None,
        _description_signature(first),
        _description_signature(second)
    ).ratio()


def _can_group_items(
    first,
    second
):
    """
    Konzervatívne všeobecné rozhodovanie.

    Položky spojíme iba keď:
    - majú rovnakú MJ,
    - nemajú konfliktné technické parametre,
    - a text je rovnaký alebo veľmi podobný.

    Pri neistote ich NEZLÚČIME.
    """

    if (
        _normalize_unit(
            first["unit"]
        )
        !=
        _normalize_unit(
            second["unit"]
        )
    ):
        return False

    first_params = (
        _extract_critical_parameters(
            first["description"]
        )
    )

    second_params = (
        _extract_critical_parameters(
            second["description"]
        )
    )

    # Ak majú oba technické parametre,
    # musia byť zhodné.
    if (
        first_params
        and second_params
        and first_params != second_params
    ):
        return False

    first_signature = (
        _description_signature(
            first["description"]
        )
    )

    second_signature = (
        _description_signature(
            second["description"]
        )
    )

    # Najbezpečnejší prípad.
    if (
        first_signature
        == second_signature
    ):
        return True

    similarity = (
        _description_similarity(
            first["description"],
            second["description"]
        )
    )

    first_code = (
        _normalize_code(
            first["code"]
        )
    )

    second_code = (
        _normalize_code(
            second["code"]
        )
    )

    # Rovnaký normalizovaný kód + veľmi podobný popis.
    # To zachytí .S / .s / bez suffixu, ale neprebije
    # rozdiel C12/15 vs C20/25, lebo ten blokuje parameter.
    if (
        first_code
        and first_code == second_code
        and similarity >= 0.78
    ):
        return True

    # Bez rovnakého kódu iba pri takmer identickom texte.
    if similarity >= 0.94:
        return True

    return False


# ==========================================================
# 3. VŠEOBECNÉ ZOSKUPOVANIE
# ==========================================================

def _make_generic_group_key(
    representative
):
    signature = (
        _description_signature(
            representative[
                "description"
            ]
        )
    )

    params = (
        _extract_critical_parameters(
            representative[
                "description"
            ]
        )
    )

    safe_signature = re.sub(
        r"[^a-z0-9]+",
        "_",
        signature
    ).strip("_")

    if len(
        safe_signature
    ) > 90:
        safe_signature = (
            safe_signature[:90]
        )

    param_text = "_".join(
        re.sub(
            r"[^a-z0-9]+",
            "_",
            param
        ).strip("_")
        for param in params
    )

    if param_text:
        return (
            safe_signature
            + "__"
            + param_text
        )

    return safe_signature


def _convert_quantity_for_ksp(
    description,
    unit,
    quantity
):
    """
    Všeobecná, konzervatívna odvodená konverzia.

    Zatiaľ prepočítava iba betónovú vrstvu:
    m2 × explicitná hrúbka = m3.

    Ak podmienky nie sú jednoznačné,
    ponechá pôvodné množstvo a MJ.
    """

    normalized = _normalize_text(
        description
    )

    normalized_unit = (
        _normalize_unit(
            unit
        )
    )

    if normalized_unit != "m2":
        return (
            quantity,
            normalized_unit
        )

    # Prevod m2 -> m3 robíme iba pri cementovom betóne,
    # keď je v popise výslovne uvedená trieda Cxx/yy.
    # Tým sa napr. asfaltový betón AC 11 O NESMIE
    # omylom prepočítať na objem.
    concrete_class = re.search(
        r"\bc\s*[0-9]+\s*/\s*[0-9]+\b",
        normalized,
        flags=re.IGNORECASE
    )

    if not concrete_class:
        return (
            quantity,
            normalized_unit
        )

    match = re.search(
        r"\bhr\.?\s*([0-9]+(?:[.,][0-9]+)?)\s*mm\b",
        normalized
    )

    if not match:
        return (
            quantity,
            normalized_unit
        )

    thickness_mm = float(
        match.group(1).replace(
            ",",
            "."
        )
    )

    return (
        quantity
        * thickness_mm
        / 1000.0,
        "m3"
    )


def aggregate_budget_items_python(items):
    """
    Všeobecné zoskupovanie bez AI.

    Algoritmus:
    1. ide položku po položke,
    2. hľadá existujúcu technicky zhodnú skupinu,
    3. ak si nie je istý, vytvorí novú skupinu,
    4. Python sčíta iba potvrdené zhody.
    """

    groups = []

    for item in items:

        matching_group = None

        for group in groups:

            if _can_group_items(
                item,
                group["representative"]
            ):
                matching_group = group
                break

        if matching_group is None:

            matching_group = {
                "representative":
                    item,
                "members":
                    []
            }

            groups.append(
                matching_group
            )

        matching_group[
            "members"
        ].append(
            item
        )

    result = []

    for group in groups:

        representative = (
            group[
                "representative"
            ]
        )

        total_quantity = 0.0
        result_unit = None
        source_rows = []

        for member in group[
            "members"
        ]:

            converted_quantity, converted_unit = (
                _convert_quantity_for_ksp(
                    member[
                        "description"
                    ],
                    member[
                        "unit"
                    ],
                    member[
                        "quantity"
                    ]
                )
            )

            # Ak by sa v skupine po odvodení objavili
            # rozdielne MJ, radšej pôvodné množstvo neprepisujeme.
            if result_unit is None:
                result_unit = (
                    converted_unit
                )

            if (
                converted_unit
                != result_unit
            ):
                converted_quantity = (
                    member[
                        "quantity"
                    ]
                )

                converted_unit = (
                    _normalize_unit(
                        member[
                            "unit"
                        ]
                    )
                )

            total_quantity += float(
                converted_quantity
            )

            source_rows.append(
                {
                    "sheet":
                        member[
                            "sheet"
                        ],
                    "row_number":
                        member[
                            "row_number"
                        ],
                    "code":
                        member[
                            "code"
                        ],
                    "description":
                        member[
                            "description"
                        ],
                    "original_quantity":
                        member[
                            "quantity"
                        ],
                    "original_unit":
                        member[
                            "unit"
                        ]
                }
            )

        params = (
            _extract_critical_parameters(
                representative[
                    "description"
                ]
            )
        )

        result.append(
            {
                "group_key":
                    _make_generic_group_key(
                        representative
                    ),
                "item_name":
                    representative[
                        "description"
                    ],
                "category":
                    "stavebna_polozka",
                "unit":
                    result_unit
                    or _normalize_unit(
                        representative[
                            "unit"
                        ]
                    ),
                "quantity":
                    round(
                        total_quantity,
                        6
                    ),
                "dimension":
                    "; ".join(
                        params
                    ),
                "material":
                    "",
                "source_rows":
                    source_rows,
                "grouping_method":
                    "python_generic"
            }
        )

    result.sort(
        key=lambda item: (
            str(
                item.get(
                    "source_rows",
                    [{}]
                )[0].get(
                    "sheet",
                    ""
                )
            ),
            int(
                item.get(
                    "source_rows",
                    [{}]
                )[0].get(
                    "row_number",
                    999999
                )
            )
        )
    )

    return result


def process_budget_python(file_bytes):
    """
    Verejná funkcia pre appku.

    - všetky detailné hárky
    - bez AI
    - všeobecné stavebné položky
    - konzervatívne zoskupovanie
    - Python sčítanie
    """

    items = extract_budget_items_python(
        file_bytes
    )

    return aggregate_budget_items_python(
        items
    )


def merge_aggregated_budget_rows(
    aggregated_lists
):
    """
    Spojí výsledky z viacerých rozpočtových Excelov.

    Znova konzervatívne:
    rovnaký group_key + rovnaká MJ.
    """

    grouped = {}

    for aggregated_rows in aggregated_lists:

        for item in aggregated_rows:

            key = (
                str(
                    item.get(
                        "group_key",
                        ""
                    )
                ),
                _normalize_unit(
                    item.get(
                        "unit",
                        ""
                    )
                )
            )

            if key not in grouped:
                grouped[key] = {
                    "group_key":
                        item.get(
                            "group_key",
                            ""
                        ),
                    "item_name":
                        item.get(
                            "item_name",
                            ""
                        ),
                    "category":
                        item.get(
                            "category",
                            ""
                        ),
                    "unit":
                        item.get(
                            "unit",
                            ""
                        ),
                    "quantity":
                        0.0,
                    "dimension":
                        item.get(
                            "dimension",
                            ""
                        ),
                    "material":
                        item.get(
                            "material",
                            ""
                        ),
                    "source_rows":
                        [],
                    "grouping_method":
                        item.get(
                            "grouping_method",
                            "python_generic"
                        )
                }

            grouped[key][
                "quantity"
            ] += float(
                item.get(
                    "quantity",
                    0
                )
                or 0
            )

            grouped[key][
                "source_rows"
            ].extend(
                item.get(
                    "source_rows",
                    []
                )
            )

    result = list(
        grouped.values()
    )

    for item in result:
        item["quantity"] = round(
            item["quantity"],
            6
        )

    return result


# ==========================================================
# SPÄTNÁ KOMPATIBILITA
# ==========================================================

def aggregate_budget_rows(classified_rows):
    """
    Staršia funkcia ostáva iba preto,
    aby prípadný starší import appku nezrútil.
    Nový rozpočet ju už nepotrebuje.
    """

    if not isinstance(
        classified_rows,
        list
    ):
        return []

    return classified_rows


# ==========================================================
# HLAVNÝ DISPEČER
# ==========================================================

def extract_text_from_file(arg1, arg2):
    """
    Podporuje obe poradia:
    extract_text_from_file(file_bytes, file_name)
    extract_text_from_file(file_name, file_bytes)
    """

    if isinstance(
        arg1,
        (bytes, bytearray)
    ):
        file_bytes = arg1
        file_name = arg2
    else:
        file_name = arg1
        file_bytes = arg2

    extension = (
        str(
            file_name
        )
        .lower()
        .split(".")[-1]
    )

    if extension == "pdf":
        return extract_text_from_pdf(
            file_bytes
        )

    if extension == "docx":
        return extract_text_from_docx(
            file_bytes
        )

    if extension in [
        "xlsx",
        "xls"
    ]:
        return extract_text_from_excel(
            file_bytes
        )

    if extension == "doc":
        raise ValueError(
            "Starý formát .doc nie je možné "
            "spoľahlivo čítať cez python-docx. "
            "Ulož dokument ako .docx."
        )

    raise ValueError(
        f"Nepodporovaný typ súboru: "
        f"{file_name}"
    )
